#-----------------------------------------------------------------------------------------------------------------------
# Name: Najiyah Williamson
# Date: 4/9/26
# Updated: 4/23/26
# BIOS 584 Cleaning Data for Project
#-----------------------------------------------------------------------------------------------------------------------

# importing pandas and loading in the data to
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document

water_pollution_disease = pd.read_csv(
    r"C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Final Project\water_pollution_disease.csv"
)

save_path = r"C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Final Project"
#-----------------------------------------------------------------------------------------------------------------------
# Cleaning the data
#-----------------------------------------------------------------------------------------------------------------------

# looking at the columns and stuff
print(water_pollution_disease)

# identifying missing values
water_pollution_disease.isna().sum()
        ### Note: Water Treatment Method has 747 missing values

# dropping missing values for water treatment method
water_pollution_disease = water_pollution_disease.dropna(subset=["Water Treatment Method"])

# double checking that they are gone
water_pollution_disease.isna().sum()

# converting years, cholera, and diarrhea from floats to integers
water_pollution_disease["Year"] = water_pollution_disease["Year"].astype(int)
water_pollution_disease["Cholera Cases per 100,000 people"] = water_pollution_disease["Cholera Cases per 100,000 people"].astype(int)
water_pollution_disease["Diarrheal Cases per 100,000 people"] = water_pollution_disease["Diarrheal Cases per 100,000 people"].astype(int)

# double checking that it worked out
print(water_pollution_disease.dtypes)

# creating a new variable to categorize lead into two risk groups
water_pollution_disease["Lead Risk"] = np.where(
    water_pollution_disease["Lead Concentration (µg/L)"] >= 15,
    "High Risk",
    "Low Risk"
)

#-------------------------------------------------------
#Subsetting the Data
#-------------------------------------------------------
# subsetting variables to the following dataset:
waterpoll_subset_USA_and_Mexico = water_pollution_disease[
    (water_pollution_disease["Country"].isin(["USA", "Mexico"])) &
    (water_pollution_disease["Year"].between(2014, 2024))
][[
    "Country",
    "Region",
    "Year",
    "Water Source Type",
    "Diarrheal Cases per 100,000 people",
    "Cholera Cases per 100,000 people",
    "Healthcare Access Index (0-100)",
    "Urbanization Rate (%)"
]]

# confirming that subsetting worked
waterpoll_subset_USA_and_Mexico.info()

# saving waterpoll_subset_USA_and_Mexico to a csv for submission
waterpoll_subset_USA_and_Mexico.to_csv(r"C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Final Project\waterpoll_subset_USA_and_Mexico.csv", index=False)

#-------------------------------------------------------

# creating a second subset of data for lead stuff
waterpoll_subset_Lead = water_pollution_disease[
    (water_pollution_disease["Year"] == 2024)
][[
    "Country",
    "Region",
    "Year",
    "Water Source Type",
    "Lead Concentration (µg/L)",
    "Lead Risk",
    "Access to Clean Water (% of Population)",
    "Healthcare Access Index (0-100)",
    "Urbanization Rate (%)"
]]

# saving waterpoll_subset_USA_and_Mexico to a csv for submission
waterpoll_subset_Lead.to_csv(r"C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Final Project\waterpoll_subset_Lead.csv", index=False)

#-----------------------------------------------------------------
# # creating third subset for all variables for US during 2024
# waterpoll_us_2024 = waterpoll_subset_USA_and_Mexico[
#     (waterpoll_subset_USA_and_Mexico["Country"] == "USA") &
#     (waterpoll_subset_USA_and_Mexico["Year"] == 2024)
# ]
#
# # confirming that this worked
# print(waterpoll_us_2024)
#
# # saving waterpoll_us_2024
# waterpoll_us_2024.to_csv(r"C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Final Project\waterpoll_us_2024.csv", index=False)

#-----------------------------------------------------------------------------------------------------------------------
# Table 1: Creating a Table 1 of USA and Mexico
#-----------------------------------------------------------------------------------------------------------------------
#-----------------------------------------------------------------
#this is for USA and Mexico Data
#-----------------------------------------------------------------
# counting the data
n_obs = len(waterpoll_subset_USA_and_Mexico)
n_countries = waterpoll_subset_USA_and_Mexico["Country"].nunique()

# function for mean (SD)
def mean_sd(series): # defining the function named mean_sd; one argument is series
    return f"{series.mean():.2f} ({series.std():.2f})" # computing mean and std of the column, rounding 2 decimals

# defining variables for the table
continuous_vars = [
    "Cholera Cases per 100,000 people",
    "Diarrheal Cases per 100,000 people",
    "Healthcare Access Index (0-100)",
    "Urbanization Rate (%)"
]

categorical_vars = [
    "Region",
    "Water Source Type"
]

# continuous table
cont_table = waterpoll_subset_USA_and_Mexico.groupby("Country")[continuous_vars].agg(mean_sd).T

# categorical table
cat_table_list = []

for var in categorical_vars: # looping over each categorical variable
    temp = pd.crosstab(waterpoll_subset_USA_and_Mexico[var], waterpoll_subset_USA_and_Mexico["Country"])
            # this code above is creating a frequency table (counts)

    # convert to n (%) while looping over each country column
    for col in temp.columns:
        total = temp[col].sum() # this is the total count per country
        temp[col] = temp[col].apply(lambda x: f"{x} ({(x/total)*100:.1f}%)") # conversion

    # label rows so adding the variable name to each category
    temp.index = [f"{var}: {i}" for i in temp.index]

    cat_table_list.append(temp) # storing this table into a list for later

cat_table = pd.concat(cat_table_list) # combine all categorical tables row-wise

# combine tables
cont_table.index = [f"{var} - Mean (SD)" for var in cont_table.index]
cont_table.index.name = "Variable" # so labeling the index column as variables
cont_table = cont_table.reset_index() # turning the index column into a normal column for csv

cat_table.index.name = "Variable"
cat_table = cat_table.reset_index()

table1USAandMexico = pd.concat([cont_table, cat_table], ignore_index=True)

print(table1USAandMexico)

table1USAandMexico.to_csv(
    fr"{save_path}\table1USAandMexico.csv",
    index=False
)

# function to create Word doc for Table 1
def save_table1USAandMexico_to_word(df, title, filename, save_path):
    doc = Document()
    doc.add_heading(title, level=1)

    # number of columns = all columns in dataframe
    n_cols = len(df.columns)

    # Create table
    table = doc.add_table(rows=1, cols=n_cols)

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])

    # Save file
    doc.save(fr"{save_path}\{filename}.docx")

save_table1USAandMexico_to_word(
    table1USAandMexico,
    title="Table 1. Descriptive Statistics (USA vs Mexico)",
    filename="table1_USA_Mexico",
    save_path=save_path
)

#-----------------------------------------------------------------------------------------------------------------------
# Table 2: Creating a Table 1 of Lead Dataset
#-----------------------------------------------------------------------------------------------------------------------
#-----------------------------------------------------------------
#this is for the lead Data
#-----------------------------------------------------------------
# defining the variables
continuous_vars = [
    "Lead Concentration (µg/L)",
    "Access to Clean Water (% of Population)",
    "Healthcare Access Index (0-100)",
    "Urbanization Rate (%)"
]

categorical_vars = [
    "Region",
    "Water Source Type"
]

# a function for mean and standard deviation
def mean_sd(x):
    return f"{x.mean():.2f} ({x.std():.2f})"

# getting a continuous table by country
cont_table = waterpoll_subset_Lead.groupby("Country")[continuous_vars].agg(mean_sd).T

# label rows
cont_table.index = [f"{var} - Mean (SD)" for var in cont_table.index]

# categorical table
cat_table_list = []

for var in categorical_vars:
    temp = pd.crosstab(waterpoll_subset_Lead[var], waterpoll_subset_Lead["Country"])

    # convert to n (%)
    for col in temp.columns:
        total = temp[col].sum()
        temp[col] = temp[col].apply(lambda x: f"{x} ({(x/total)*100:.1f}%)")

    # label rows
    temp.index = [f"{var}: {i}" for i in temp.index]

    cat_table_list.append(temp)

cat_table = pd.concat(cat_table_list)

# combining the tables
cont_table.index.name = "Variable"
cont_table = cont_table.reset_index()

cat_table.index.name = "Variable"
cat_table = cat_table.reset_index()

table1_Lead = pd.concat([cont_table, cat_table], ignore_index=True)

print(table1_Lead)

table1_Lead.to_csv(
    r"C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Final Project\table1_Lead_2024.csv",
    index=False
)

# creating table 1
# function to create Word doc for Table 1
def save_table1Lead_to_word(df, title, filename, save_path):
    doc = Document()
    doc.add_heading(title, level=1)

    # number of columns = all columns in dataframe
    n_cols = len(df.columns)

    # Create table
    table = doc.add_table(rows=1, cols=n_cols)

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])

    # Save file
    doc.save(fr"{save_path}\{filename}.docx")

save_table1Lead_to_word(
    table1_Lead,
    title="Table 1. Descriptive Statistics for Lead Exposure (2024)",
    filename="table1_Lead_2024",
    save_path=save_path
)

#-----------------------------------------------------------------
# Visual 1: Healthcare Access Index for USA and Mexico (Histogram)
#-----------------------------------------------------------------
plt.figure(figsize=(6,5))

sns.histplot(
    data=waterpoll_subset_USA_and_Mexico,
    x="Healthcare Access Index (0-100)",
    hue="Country",
    bins=20,
    multiple="dodge"
)

plt.title("Healthcare Access Index Distribution by Country")
plt.xlabel("Healthcare Access Index (0-100)")
plt.ylabel("Count")

plt.savefig(
    fr"{save_path}\healthcare_access_hist_country.png",
    dpi=300,
    bbox_inches="tight"
)

plt.close()

#-----------------------------------------------------------------
# Visual 2: Healthcare Access Index for USA and Mexico (Boxplot)
#-----------------------------------------------------------------
plt.figure(figsize=(6,5))

sns.boxplot(
    data=waterpoll_subset_USA_and_Mexico,
    x="Country",
    y="Healthcare Access Index (0-100)"
)

plt.title("Healthcare Access Index by Country")
plt.xlabel("Country")
plt.ylabel("Healthcare Access Index")

plt.savefig(
    fr"{save_path}\healthcare_access_boxplot.png",
    dpi=300,
    bbox_inches="tight"
)

plt.close()





# #-----------------------------------------------------------------------------------------------------------------------
# # Table 3: Creating a table of US region summary statistics
# #-----------------------------------------------------------------------------------------------------------------------
# # for variables that are below the threshold, so low healthcare access index in the US during 2024
#
# # defining "low healthcare access"
# # bottom 25% of Healthcare Access Index
# threshold = waterpoll_us_2024["Healthcare Access Index (0-100)"].quantile(0.25)
# low_access = waterpoll_us_2024[
#     waterpoll_us_2024["Healthcare Access Index (0-100)"] <= threshold
#     ].copy()
#
# # helper function for mean(SD)
# def mean_sd(series):
#     return f"{series.mean():.2f} ({series.std():.2f})"
#
# # group by Region and compute descriptive statistics
# table2_us_lowaccess = pd.DataFrame({
#     "Variable": [
#         "Country of interest",
#         "Year of analysis",
#         "Number of observations",
#         "Number of regions",
#         "Healthcare Access Index (0–100)",
#         "Cholera incidence (per 100,000)",
#         "Diarrheal incidence (per 100,000)",
#         "Access to clean water (%)",
#         "Sanitation coverage (%)"
#     ],
#     "Mean (SD)": [
#         "United States",                     # country
#         "2024",                               # year
#         len(low_access),                      # n
#         low_access["Region"].nunique(),       # number of regions
#         mean_sd(low_access["Healthcare Access Index (0-100)"]),
#         mean_sd(low_access["Cholera Cases per 100,000 people"]),
#         mean_sd(low_access["Diarrheal Cases per 100,000 people"]),
#         mean_sd(low_access["Access to Clean Water (% of Population)"]),
#         mean_sd(low_access["Sanitation Coverage (% of Population)"])
#     ]
# })
#
#
# table2_us_lowaccess
#
# table2_us_lowaccess.to_csv(
#     fr"{save_path}\table2_low_access_2024.csv",
#     index=False
# )
#
# save_table1_to_word(
#     table2_us_lowaccess,
#     title="Table 2. Low Healthcare Access Regions (U.S., 2024)",
#     filename="table2_us_lowaccess",
#     save_path=save_path
# )
#
# #-----------------------------------------------------------------------------------------------------------------------
# # Visual 1: Disease incidence over time
# #-----------------------------------------------------------------------------------------------------------------------
# # turning to long format
# subset_long = waterpoll_subset.melt(
#     id_vars=["Country", "Year"],
#     value_vars=[
#         "Cholera Cases per 100,000 people",
#         "Diarrheal Cases per 100,000 people"
#     ],
#     var_name="Disease",
#     value_name="Incidence"
# )
#
# # plotting disease incidence
# plt.figure(figsize=(10, 6))
#
# sns.lineplot(
#     data=subset_long,
#     x="Year",
#     y="Incidence",
#     hue="Disease",
#     estimator="mean",
#     errorbar="sd",
#     linewidth=2
# )
#
# plt.title("Trends in Cholera and Diarrheal Disease Incidence Over Time")
# plt.ylabel("Incidence per 100,000 people")
# plt.xlabel("Year")
# plt.tight_layout()
# plt.show()
#
# plt.savefig(
#     fr"{save_path}\disease_trends.png",
#     dpi=300,
#     bbox_inches="tight"
# )
# plt.close()