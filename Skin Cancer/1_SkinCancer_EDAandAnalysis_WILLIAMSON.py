import os
import sys

import seaborn as sns
import matplotlib.pyplot as plt
import numpy as np
import scipy.stats as stats
import pandas as pd
from scipy.stats import chi2_contingency


print(os.getcwd()) #get current working directory
sys.path.append(r'C:\Users\Najiy\OneDrive\Desktop\PycharmProjects\BIOS584-Python\Skin Cancer')
from SkinCancer_Analysis_Functions import chisq_res

#Read in cleaned data
skc_cc = pd.read_excel("cleaned_skc_cc.xlsx")
skc = pd.read_excel("cleaned_skc.xlsx")
#----------------------------------------------------------------------------
# Research Question 1
#----------------------------------------------------------------------------
# conversion of num_primary_malignancies
skc_cc["num_primary_malignancies"] = skc_cc["num_primary_malignancies"].replace({
    "One cancer": 1,
    "More than one cancer": 2
})

print(skc_cc["num_primary_malignancies"].describe())

#---------------------------------------------------------------------
# T-test with side by side boxplots
#---------------------------------------------------------------------------
sns.boxplot(x='Race_binary', y='num_primary_malignancies', data=skc_cc)
plt.title("Number of Primary Malignancies by Race")
plt.show()

#-------------------------------------------------------------------------------
# Summary statistics
#-------------------------------------------------------------------------------
summary = skc_cc.groupby("Race_binary")["num_primary_malignancies"].agg(
    mean="mean",
    std="std",
    median="median",
    maximum="max"
)
print(summary)

#-------------------------------------------------------------------------------
# T-test
#-------------------------------------------------------------------------------
poc = skc_cc[skc_cc["Race_binary"] == "POC"]["num_primary_malignancies"]
white = skc_cc[skc_cc["Race_binary"] == "White"]["num_primary_malignancies"]

# compute summary stats
mean_poc, sd_poc = poc.mean(), poc.std()
median_poc, max_poc = poc.median(), poc.max()

mean_white, sd_white = white.mean(), white.std()
median_white, max_white = white.median(), white.max()

# t-test
tstat, pval = stats.ttest_ind(white, poc, equal_var=False)

# assumption check
stats.shapiro(poc)
stats.shapiro(white)

# difference in means + 95% CI
diff = mean_white - mean_poc
se = np.sqrt(white.var(ddof=1)/len(white) + poc.var(ddof=1)/len(poc))

ci_low = diff - 1.96 * se
ci_high = diff + 1.96 * se

table = pd.DataFrame({
    "Number of Primary Malignancies": ["POC patients", "White patients"],
    "Mean (SD)": [
        f"{mean_poc:.2f} ({sd_poc:.2f})",
        f"{mean_white:.2f} ({sd_white:.2f})"
    ],
    "Median": [median_poc, median_white],
    "Maximum": [max_poc, max_white],
    "Difference in Means [95% CI]": [f"{diff:.2f} [{ci_low:.2f}, {ci_high:.2f}]", ""],
    "P-value": [f"{pval:.4f}", ""]
})

print(table)

table.to_csv("rq1_table.csv", index=False)

#----------------------------------------------------------------------------
# Research Question 2
#----------------------------------------------------------------------------
#-------------------------------
# Chi-square Test 1
# Race vs M1_Location_risk
#-------------------------------
chi2_1, p1, dof1, expected1, table1 = chisq_res(
    "Race_binary",
    "M1_Location_risk",
    skc
)

print("\nTest 1 Results")
print("Chi-square:", chi2_1)
print("p-value:", p1)
print("\nExpected Counts:")
print(pd.DataFrame(expected1, index=table1.index, columns=table1.columns).round(2))

#-------------------------------
# Chi-square Test 2
# Race vs cancer_type
#-------------------------------
chi2_2, p2, dof2, expected2, table2 = chisq_res(
    "Race_binary",
    "cancer_type",
    skc
)

print("\nTest 2 Results")
print("Chi-square:", chi2_2)
print("p-value:", p2)
print("\nExpected Counts:")
print(pd.DataFrame(expected2, index=table2.index, columns=table2.columns).round(2))

# checking the info
skc['Race_binary'].value_counts()
skc['M1_Location_risk'].value_counts()
skc['cancer_type'].value_counts()

# making the really large table
def chisq_table(df, row_var, col_var):

    # contingency table (counts)
    table = pd.crosstab(df[row_var], df[col_var])

    # chi-square test
    chi2, p, dof, expected = chi2_contingency(table)

    # convert to row percentages
    row_pct = table.div(table.sum(axis=1), axis=0) * 100

    return table, row_pct, chi2, p

# -------------------------
# Test 1: Tumor location risk
# -------------------------
table1, pct1, chi2_1, p1 = chisq_table(
    skc,
    "Race_binary",
    "M1_Location_risk"
)

# -------------------------
# Test 2: Cancer type
# -------------------------
table2, pct2, chi2_2, p2 = chisq_table(
    skc,
    "Race_binary",
    "cancer_type"
)

# sample sizes
total_n = len(skc)
poc_n = skc[skc["Race_binary"] == "POC"].shape[0]
white_n = skc[skc["Race_binary"] == "White"].shape[0]
print(total_n, poc_n, white_n)

# adding a total column
total_counts1 = table1.sum(axis=0) # axis=0 means combining things vertically
total_pct1 = (total_counts1 / total_counts1.sum()) * 100 # getting percentages

total_counts2 = table2.sum(axis=0)
total_pct2 = (total_counts2 / total_counts2.sum()) * 100 # getting percentages

# tumor table
tumor_table = pd.DataFrame({
    "Category": table1.columns,
    "Total n": total_counts1.values,
    "Total %": total_pct1.round(1).values,
    "POC n": table1.loc["POC"].values,
    "POC %": pct1.loc["POC"].round(1).values,
    "White n": table1.loc["White"].values,
    "White %": pct1.loc["White"].round(1).values,
})

tumor_table["χ²"] = ""
tumor_table["p-value"] = ""

# cancer table
cancer_table = pd.DataFrame({
    "Category": table2.columns,
    "Total n": total_counts2.values,
    "Total %": total_pct2.round(1).values,
    "POC n": table2.loc["POC"].values,
    "POC %": pct2.loc["POC"].round(1).values,
    "White n": table2.loc["White"].values,
    "White %": pct2.loc["White"].round(1).values,
})

cancer_table["χ²"] = ""
cancer_table["p-value"] = ""

# trying to add section headers
tumor_header = pd.DataFrame({
    "Category": ["Tumor Location Risk"],
    "Total n": [""], "Total %": [""],
    "POC n": [""], "POC %": [""],
    "White n": [""], "White %": [""],
    "χ²": [round(chi2_1, 2)],
    "p-value": [round(p1, 4)]
})

cancer_header = pd.DataFrame({
    "Category": ["Cancer Type"],
    "Total n": [""], "Total %": [""],
    "POC n": [""], "POC %": [""],
    "White n": [""], "White %": [""],
    "χ²": [round(chi2_2, 2)],
    "p-value": [round(p2, 4)]
})

# combining into one table
final_table = pd.concat([
    tumor_header,
    tumor_table,
    cancer_header,
    cancer_table
], ignore_index=True) # index is setting the row numbers to cleaner values that are in order

print(final_table)

# exporting to word
from docx import Document

doc = Document()
doc.add_heading('Table 2. Chi-Square Analyses to Examine Race Differences', 1)

word_table = doc.add_table(rows=1, cols=len(final_table.columns))

# header
for i, col in enumerate(final_table.columns):
    word_table.rows[0].cells[i].text = col

# rows
for _, row in final_table.iterrows():
    cells = word_table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(val)

doc.save("Table2_ChiSquare.docx")